# =========================================================
# --- ФІНАЛЬНА ВЕРСІЯ (З ВИПРАВЛЕНИМИ ПРОКРУТКАМИ) ---
# =========================================================

import sys
import os
import time 
import shutil 
import torch 
import pathlib 
import tempfile 
import re 
import io 
import tkinter as tk

# --- 0. НАЛАШТУВАННЯ ШЛЯХІВ ---
# ВАЖЛИВО: Переконайтеся, що ці шляхи коректні на вашій системі!
WHISPER_CACHE_ROOT = r"C:\Whisper_cache"
os.environ["XDG_CACHE_HOME"] = WHISPER_CACHE_ROOT
SITE_PACKAGES_PATH = r"C:\Users\Tata\AppData\Local\Programs\Python\Python311\Lib\site-packages"

if SITE_PACKAGES_PATH not in sys.path and os.path.exists(SITE_PACKAGES_PATH):
    sys.path.append(SITE_PACKAGES_PATH)
    print(f"✅ [sys.path] SITE-PACKAGES ПРИМУСОВО додано: {SITE_PACKAGES_PATH}")

try:
    SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
    if SCRIPT_DIR not in sys.path:
        sys.path.append(SCRIPT_DIR)
        print(f"✅ [sys.path] Папка скрипта додана: {SCRIPT_DIR}")
except NameError:
    pass 

FFMPEG_BIN_PATH = r"C:\Users\Tata\ffmpeg-2025-09-15-git-16b8a7805b-full_build\bin" 

if os.path.exists(FFMPEG_BIN_PATH):
    os.environ["PATH"] += os.pathsep + FFMPEG_BIN_PATH
    print(f"✅ [ENV] FFMPEG шлях додано: {FFMPEG_BIN_PATH}")
else:
    print(f"❌ FFMPEG Попередження: Не знайдено за адресою: {FFMPEG_BIN_PATH}. Перевірте шлях!")

# --- 3. Імпорт необхідних бібліотек ---
try:
    import numpy as np
    import whisper
    from pydub import AudioSegment 
    import matplotlib.pyplot as plt
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    import zipfile
    from googletrans import Translator 
    from gtts import gTTS 
    from tkinter import filedialog, messagebox, scrolledtext
    import customtkinter as ctk 
except ImportError as e:
    print(f"❌ КРИТИЧНА ПОМИЛКА: Не вдалося імпортувати одну з бібліотек. Встановіть її: {e}")
    if "gTTS" in str(e):
         print(">>> Для генерації TTS-аудіо виконайте в консолі: pip install gTTS")
    sys.exit(1)

DEVICE = "cuda" if torch.cuda.is_available() else "cpu"

# =========================================================
# --- КЛАС ДОДАТКУ ---
# =========================================================

class ConsecutiveTrainerApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("🎧 Consecutive Translation Trainer")
        # БЕЗПЕЧНА ВИСОТА ТА ЗМІНА РОЗМІРУ
        self.geometry("650x850") 
        self.resizable(True, True) 
        ctk.set_appearance_mode("System")
        ctk.set_default_color_theme("blue")

        self.audio_filepath = None
        self.model_name = ctk.StringVar(value="small") 
        self.pause_duration = ctk.IntVar(value=7)       
        self.output_dir = ctk.StringVar(value=os.path.join(os.path.expanduser("~"), "Desktop")) 
        self.source_lang = ctk.StringVar(value="uk") 
        self.target_lang = ctk.StringVar(value="en") 
        self.audio_type = ctk.StringVar(value="Human (Default)")
        self.output_format = ctk.StringVar(value="MP3") 
        self.generate_tts = ctk.BooleanVar(value=False) 
        self.segmentation_choice = ctk.StringVar(value="Refined (Full Sentences)") 
        self.note_template = ctk.StringVar(value="Symbolic/Keyword (2 Колонки)")
        
        self.device = DEVICE
        self.translator = Translator() 
        
        self.AUDIO_END_BUFFER_MS = 600       
        self.AUDIO_START_BUFFER_MS = 400     
        self.FADE_DURATION_MS = 200          
        self.MIN_SEGMENT_DURATION_MS = 1000 

        self._build_ui()
        self.log_message(f"✅ Система ініціалізована. Пристрій: {self.device.upper()}")
        self.log_message(f"✅ Кеш моделі встановлено у: {WHISPER_CACHE_ROOT}\\whisper\\")
        self.log_message("Оберіть файл та натисніть 'ЗАПУСТИТИ ОБРОБКУ'.")

    # --- UI та Логгінг УТИЛІТИ ---
    
    def _autoscroll_y(self, *args):
        """Гарантує, що скролбар логу правильно відображається."""
        self.log_text.yview(*args)

    def select_audio_file(self):
        filepath = filedialog.askopenfilename(filetypes=[("Audio Files", "*.mp3 *.wav *.ogg *.flac")])
        if filepath:
            self.audio_filepath = filepath
            self.file_label.configure(text=os.path.basename(filepath))
            self.log_message(f"✅ Файл вибрано: {os.path.basename(filepath)}")

    def select_output_dir(self):
        dirpath = filedialog.askdirectory()
        if dirpath:
            self.output_dir.set(dirpath)
            self.log_message(f"✅ Папка збереження встановлена: {dirpath}")
            
    def log_message(self, message):
        """Вставляє повідомлення в лог та автоматично прокручує його до кінця."""
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END) # Гарантує, що останні повідомлення видно
        self.update_idletasks()

    def set_running_state(self, is_running):
        if is_running:
            self.start_button.configure(text="ОБРОБКА...", state="disabled", fg_color="gray")
        else:
            self.start_button.configure(text="10. ЗАПУСТИТИ ОБРОБКУ", state="normal", fg_color="green")

    def _build_ui(self):
        
        # 1. ОСНОВНИЙ КОНТЕЙНЕР ДЛЯ НАЛАШТУВАНЬ З ПРОКРУТКОЮ
        scrollable_settings_frame = ctk.CTkScrollableFrame(self, label_text="Налаштування Обробки (Scrollable)", 
                                                             scrollbar_fg_color="transparent")
        # fill="x" дозволяє логу отримати залишковий вертикальний простір.
        scrollable_settings_frame.pack(padx=20, pady=(10, 5), fill="x") 

        # 2. ФРЕЙМ НАЛАШТУВАНЬ УСЕРЕДИНІ КОНТЕЙНЕРА
        settings_frame = ctk.CTkFrame(scrollable_settings_frame)
        settings_frame.pack(padx=0, pady=0, fill="x", expand=True) 
        settings_frame.grid_columnconfigure(0, weight=3)
        settings_frame.grid_columnconfigure(1, weight=1)

        # 1. ФАЙЛ (рядки 0, 1)
        ctk.CTkLabel(settings_frame, text="1. Вхідний Аудіофайл:", font=("Arial", 14, "bold")).grid(row=0, column=0, padx=10, pady=5, sticky="w", columnspan=2)
        self.file_label = ctk.CTkLabel(settings_frame, text="Файл не вибрано.", text_color="gray", wraplength=400)
        self.file_label.grid(row=1, column=0, padx=10, pady=5, sticky="w")
        ctk.CTkButton(settings_frame, text="Вибрати MP3/WAV", command=self.select_audio_file).grid(row=1, column=1, padx=10, pady=5)
        
        # 2. ПАПКА (рядки 2, 3)
        ctk.CTkLabel(settings_frame, text="2. Папка для Збереження:", font=("Arial", 14, "bold")).grid(row=2, column=0, padx=10, pady=5, sticky="w", columnspan=2)
        self.output_label = ctk.CTkLabel(settings_frame, textvariable=self.output_dir, text_color="green", anchor="w", wraplength=400)
        self.output_label.grid(row=3, column=0, padx=10, pady=5, sticky="ew")
        ctk.CTkButton(settings_frame, text="Вибрати Папку", command=self.select_output_dir).grid(row=3, column=1, padx=10, pady=5)
        
        # 3. МОДЕЛЬ/ПАУЗА (рядки 4, 5, 6)
        ctk.CTkLabel(settings_frame, text="3. Налаштування Моделі та Паузи:", font=("Arial", 14, "bold")).grid(row=4, column=0, padx=10, pady=5, sticky="w", columnspan=2)
        
        ctk.CTkLabel(settings_frame, text="Модель Whisper:").grid(row=5, column=0, padx=10, pady=2, sticky="w")
        model_options = ["tiny", "base", "small", "medium", "large"]
        ctk.CTkOptionMenu(settings_frame, variable=self.model_name, values=model_options).grid(row=5, column=1, padx=10, pady=2, sticky="ew")

        ctk.CTkLabel(settings_frame, text="Тривалість паузи (сек):").grid(row=6, column=0, padx=10, pady=2, sticky="w")
        ctk.CTkEntry(settings_frame, textvariable=self.pause_duration, width=80).grid(row=6, column=1, padx=10, pady=2, sticky="w")
        
        # 4. МОВИ (рядки 7, 8, 9)
        ctk.CTkLabel(settings_frame, text="4. Налаштування Мов (ISO 639-1):", font=("Arial", 14, "bold")).grid(row=7, column=0, padx=10, pady=5, sticky="w", columnspan=2)
        
        lang_options = ["uk", "en", "ru", "de", "fr", "es", "auto"] 
        
        ctk.CTkLabel(settings_frame, text="Мова Аудіо (Транскрипція):").grid(row=8, column=0, padx=10, pady=2, sticky="w")
        ctk.CTkOptionMenu(settings_frame, variable=self.source_lang, values=lang_options).grid(row=8, column=1, padx=10, pady=2, sticky="ew")

        ctk.CTkLabel(settings_frame, text="Мова Перекладу (Ключ):").grid(row=9, column=0, padx=10, pady=2, sticky="w")
        ctk.CTkOptionMenu(settings_frame, variable=self.target_lang, values=lang_options[:-1]).grid(row=9, column=1, padx=10, pady=2, sticky="ew")
        
        # 5. ОБРОБКА ОРИГІНАЛУ (рядки 10, 11)
        ctk.CTkLabel(settings_frame, text="5. Обробка Оригінального Аудіо:", font=("Arial", 14, "bold")).grid(row=10, column=0, padx=10, pady=5, sticky="w", columnspan=2)
        audio_type_options = ["Human (Default)", "TTS (Apply Fade-in/out)"]
        ctk.CTkOptionMenu(settings_frame, variable=self.audio_type, values=audio_type_options).grid(row=11, column=0, padx=10, pady=2, sticky="ew", columnspan=2)

        # 6. ГЕНЕРАЦІЯ TTS (рядок 12)
        ctk.CTkLabel(settings_frame, text="6. Створити Додаткове TTS Аудіо:", font=("Arial", 14, "bold")).grid(row=12, column=0, padx=10, pady=5, sticky="w")
        ctk.CTkCheckBox(settings_frame, text="Генерувати TTS-версію", variable=self.generate_tts).grid(row=12, column=1, padx=10, pady=5, sticky="ew")
        
        # 7. СЕГМЕНТАЦІЯ (рядок 13)
        ctk.CTkLabel(settings_frame, text="7. Логіка Сегментації Нотаток:", font=("Arial", 14, "bold")).grid(row=13, column=0, padx=10, pady=5, sticky="w")
        segmentation_options = [
            "Refined (Full Sentences)", 
            "Whisper (Default, Shorter Segments)"
        ]
        ctk.CTkOptionMenu(settings_frame, variable=self.segmentation_choice, values=segmentation_options).grid(row=13, column=1, padx=10, pady=5, sticky="ew")
        
        # 8. ШАБЛОН НОТУВАННЯ (рядок 14) 
        ctk.CTkLabel(settings_frame, text="8. Шаблон Нотування (DOCX):", font=("Arial", 14, "bold")).grid(row=14, column=0, padx=10, pady=5, sticky="w")
        template_options = [
            "Handout (Текст та Ключ)",
            "Symbolic/Keyword (2 Колонки)",
            "Усі 5 Шаблонів (Окремі файли)"
        ]
        ctk.CTkOptionMenu(settings_frame, variable=self.note_template, values=template_options).grid(row=14, column=1, padx=10, pady=5, sticky="ew")
        
        # 9. ФОРМАТ (рядки 15, 16)
        ctk.CTkLabel(settings_frame, text="9. Формат Вихідного Аудіо:", font=("Arial", 14, "bold")).grid(row=15, column=0, padx=10, pady=5, sticky="w", columnspan=2)
        output_format_options = ["MP3", "WAV", "FLAC"]
        ctk.CTkOptionMenu(settings_frame, variable=self.output_format, values=output_format_options).grid(row=16, column=0, padx=10, pady=2, sticky="ew", columnspan=2)


        # 3. КНОПКА ЗАПУСКУ
        self.start_button = ctk.CTkButton(self, text="10. ЗАПУСТИТИ ОБРОБКУ", command=self.start_processing, height=40, font=("Arial", 16, "bold"), fg_color="green")
        self.start_button.pack(padx=20, pady=10, fill="x")

        # 4. ЛОГ (використовує expand=True для заповнення залишкового простору)
        log_frame = ctk.CTkFrame(self)
        log_frame.pack(padx=20, pady=(0, 10), fill="both", expand=True)

        ctk.CTkLabel(log_frame, text="Лог Виконання:", font=("Arial", 14, "bold")).pack(padx=10, pady=(10, 5), anchor="w")
        
        # ScrolledText: Вбудований механізм прокрутки
        self.log_text = scrolledtext.ScrolledText(log_frame, 
                                                  wrap=tk.WORD, 
                                                  height=5, # Мінімальна висота 5 рядків
                                                  bg="#343638", 
                                                  fg="white", 
                                                  bd=0, 
                                                  relief="flat", 
                                                  font=("Consolas", 10),
                                                  yscrollcommand=self._autoscroll_y) # <-- Виправлення прокрутки
        
        self.log_text.pack(padx=10, pady=(0, 10), fill="both", expand=True)
        self.log_text.bind("<KeyPress>", lambda e: "break") 

    def start_processing(self):
        if not self.audio_filepath or not os.path.exists(self.audio_filepath):
            messagebox.showerror("Помилка", "Будь ласка, виберіть аудіофайл.")
            return

        try:
            pause_seconds = self.pause_duration.get()
            if not isinstance(pause_seconds, int) or pause_seconds <= 0:
                messagebox.showerror("Помилка", "Тривалість паузи має бути цілим числом, більшим за 0.")
                return

            self.log_text.delete(1.0, tk.END) 
            self.set_running_state(True)
            self.log_message("=========================================")
            self.log_message(f"🤖 Починаю обробку. Модель: {self.model_name.get()}, Пауза: {pause_seconds}с")
            self.log_message(f"📝 Мова транскрипції: {self.source_lang.get().upper()}. Мова перекладу: {self.target_lang.get().upper()}")
            self.log_message(f"🔊 Обробка оригіналу: {self.audio_type.get()}")
            self.log_message(f"🗣️ Генерація TTS-версії: {'ТАК' if self.generate_tts.get() else 'НІ'}")
            self.log_message(f"✨ Логіка сегментації: {self.segmentation_choice.get()}")
            self.log_message(f"✍️ Шаблон нотування: {self.note_template.get()}")
            self.log_message(f"💾 Вихідний формат: {self.output_format.get()}") 
            self.log_message("=========================================")

            import threading
            thread = threading.Thread(target=self._process_in_thread, daemon=True)
            thread.start()

        except Exception as e:
            messagebox.showerror("Помилка", f"Помилка в налаштуваннях: {e}")
            self.log_message(f"❌ Критична помилка в налаштуваннях: {e}")
            self.set_running_state(False)

    def _process_in_thread(self):
        temp_dir = None
        try:
            temp_dir = self._process_audio()
        except Exception as e:
            self.log_message(f"❌ Критична помилка під час обробки: {e}")
            if self.winfo_exists():
                self.after(0, lambda: messagebox.showerror("Критична Помилка", f"Обробка не вдалася. Деталі у лозі. Помилка: {e}"))
        finally:
            if self.winfo_exists():
                self.after(0, lambda: self.set_running_state(False))
                if temp_dir:
                    self.after(0, lambda: self._cleanup(temp_dir))
                self.log_message("\n✨ Обробка завершена.")
    
    def _refine_segments_by_sentences(self, segments):
        SENTENCE_ENDINGS = ['.', '?', '!', '...']
        all_words = []
        for segment in segments:
            if segment.get('words'):
                all_words.extend(segment['words'])
        
        if not all_words: return [] 

        refined_segments = []
        current_segment_words = []
        
        for word_data in all_words:
            current_segment_words.append(word_data)
            word = word_data['word']
            
            ends_with_punctuation = any(word.strip().endswith(p) for p in SENTENCE_ENDINGS)
            segment_text_so_far = "".join([w['word'] for w in current_segment_words]).strip()
            text_ends_with_punctuation = any(segment_text_so_far.endswith(p) for p in SENTENCE_ENDINGS)
            
            if (ends_with_punctuation or text_ends_with_punctuation):
                start = current_segment_words[0]['start']
                end = current_segment_words[-1]['end']
                text = "".join([w['word'] for w in current_segment_words]).strip()
                text = re.sub(r'\s+([?.!])', r'\1', text) 

                if end > start and text:
                    refined_segments.append({'start': start, 'end': end, 'text': text})
                
                current_segment_words = [] 
        
        if current_segment_words:
            start = current_segment_words[0]['start']
            end = current_segment_words[-1]['end']
            text = "".join([w['word'] for w in current_segment_words]).strip()
            text = re.sub(r'\s+([?.!])', r'\1', text)

            if end > start and text:
                refined_segments.append({'start': start, 'end': end, 'text': text})
                
        return refined_segments

    # --- Функції для DOCX ---
    def _add_document_header(self, doc, base_name, doc_info):
        doc.add_heading(f"Матеріали для тренування: {base_name}", level=1)
        
        table = doc.add_table(rows=len(doc_info) // 2 + 1, cols=4, style='Table Grid')
        table.autofit = False
        table.columns[0].width = Inches(1.8)
        table.columns[1].width = Inches(1.5)
        table.columns[2].width = Inches(1.8)
        table.columns[3].width = Inches(1.5)
        
        table.cell(0, 0).text = 'Налаштування:'
        table.cell(0, 0).paragraphs[0].runs[0].bold = True
        
        items = list(doc_info.items())
        for i, (key, value) in enumerate(items):
            row_index = i // 2
            col_index = (i % 2) * 2
            
            if row_index == 0 and col_index == 0: continue
            
            row = table.rows[row_index + 1].cells if i >= 2 else table.rows[row_index].cells
            row[col_index].text = key + ':'
            row[col_index + 1].text = str(value)

    def _create_handout_layout(self, document, text_segments):
        """Створює роздатковий матеріал (текст + ключ)."""
        document.add_heading("Шаблон 1/5: Handout (Текст та Ключ)", level=2)
        document.add_paragraph("Простий роздатковий матеріал: оригінал та повний переклад-ключ. Ідеально для фінальної перевірки.")
        document.add_paragraph("---")
        
        for i, segment in enumerate(text_segments):
            p = document.add_paragraph()
            p.add_run(f"[{i+1}]. Оригінал: ").bold = True
            p.add_run(f" {segment['text']}\n")
            p.add_run("   Переклад (Ключ): ").bold = True
            p.add_run(segment['translated_text']).italic = True
            document.add_paragraph("---", style='Normal')


    def _create_vertical_layout(self, document, text_segments, target_lang):
        """Створює єдину велику таблицю Vertical (3 Columns / Rozan)."""
        document.add_heading("Шаблон 2/5: Вертикальне Нотування (3 Колонки / Розан)", level=2)
        document.add_paragraph("Цей шаблон використовує 3 вертикальні колонки: Оригінал/Таймінг, Нотатки, Переклад (Ключ).")
        document.add_paragraph("---")
        
        table = document.add_table(rows=1, cols=3, style='Table Grid')
        table.autofit = False
        table.allow_autofit = False
        
        # Ширина: 1.0 (Index/Time) + 4.0 (Notes) + 1.5 (Key) = 6.5 дюймів
        table.columns[0].width = Inches(1.0) 
        table.columns[1].width = Inches(4.0) 
        table.columns[2].width = Inches(1.5) 
        
        row_header = table.rows[0].cells
        row_header[0].text = 'Index / Time' 
        row_header[1].text = 'Notes / Symbols' 
        row_header[2].text = f'Key ({target_lang.upper()})' 
        
        for i, segment in enumerate(text_segments):
            row = table.add_row().cells
            p1 = row[0].paragraphs[0]
            p1.add_run(f"[{i+1}]").bold = True
            p1.add_run(f" ({(segment['start']):.1f}s - {segment['end']:.1f}s)")
            row[0].add_paragraph(segment['text'])
            
            row[1].text = '' 
            row[2].text = segment['translated_text']
            
            try:
                row.height = Inches(1.5) 
            except Exception:
                pass 

        document.add_paragraph("---", style='Normal')


    def _create_horizontal_layout(self, document, text_segments, target_lang):
        """Створює єдину велику таблицю Horizontal (Main Notes)."""
        document.add_heading("Шаблон 3/5: Горизонтальне Нотування (Лінійні Нотатки)", level=2)
        document.add_paragraph("Цей шаблон надає велику область для лінійного нотування під кожним сегментом.")
        document.add_paragraph("---")
        
        for i, segment in enumerate(text_segments):
            p_info = document.add_paragraph()
            p_info.add_run(f"[{i+1}]. Оригінал: ").bold = True
            p_info.add_run(f" {segment['text']}\n")
            p_info.add_run("   Переклад (Ключ): ").bold = True
            p_info.add_run(segment['translated_text']).italic = True
            
            table = document.add_table(rows=1, cols=1, style='Table Grid')
            table.autofit = False
            table.allow_autofit = False
            
            # Ширина: 6.5 дюймів
            table.columns[0].width = Inches(6.5) 
            
            cell = table.cell(0, 0)
            cell.text = 'Notes:'
            
            row = table.add_row()
            row.cells[0].text = '' 
            try:
                row.height = Inches(1.5)
            except Exception:
                pass
            
            document.add_paragraph("---", style='Normal')

    def _create_symbolic_keyword_layout(self, document, text_segments, target_lang):
        """Створює єдину велику таблицю для нотування ключовими словами та символами (2 колонки)."""
        document.add_heading("Шаблон 4/5: Нотування Ключовими Словами та Символами (2 Колонки)", level=2)
        document.add_paragraph("Цей шаблон розділяє область для нотування на Ключові Слова/Символи та Основні Нотатки.")
        document.add_paragraph("---")
        
        for i, segment in enumerate(text_segments):
            p_info = document.add_paragraph()
            p_info.add_run(f"[{i+1}]. Оригінал: ").bold = True
            p_info.add_run(f" {segment['text']}\n")
            p_info.add_run("   Переклад (Ключ): ").bold = True
            p_info.add_run(segment['translated_text']).italic = True

            table = document.add_table(rows=1, cols=2, style='Table Grid')
            table.autofit = False
            table.allow_autofit = False
            
            # Ширина: 2.0 (Keywords/Symbols) + 4.5 (Main Notes) = 6.5 дюймів
            table.columns[0].width = Inches(2.0) 
            table.columns[1].width = Inches(4.5) 
            
            row = table.row_cells(0)
            row[0].text = 'Keywords/Symbols'
            row[1].text = 'Main Notes'
            
            row = table.add_row()
            row.cells[0].text = '' 
            row.cells[1].text = ''
            try:
                row.height = Inches(1.5)
            except Exception:
                pass
            
            document.add_paragraph("---", style='Normal')

    def _create_cornell_layout(self, document, text_segments, target_lang):
        """Створює єдину велику таблицю Корнельської системи для всіх сегментів."""
        document.add_heading("Шаблон 5/5: Корнельська Система (Cornell Notes)", level=2)
        document.add_paragraph("Велика Корнельська таблиця: ліва колонка - Питання/Ключові слова, права - Основні нотатки.")
        document.add_paragraph("---")
        
        table = document.add_table(rows=1, cols=2, style='Table Grid')
        table.autofit = False
        table.allow_autofit = False
        
        # Ширина: 2.0 (Cues) + 4.5 (Notes Area) = 6.5 дюймів
        table.columns[0].width = Inches(2.0) 
        table.columns[1].width = Inches(4.5) 
        
        row_header = table.row_cells(0)
        row_header[0].text = 'Cues / Questions:'
        row_header[1].text = 'Notes Area (Original Text Segments and Key Translations):'
        
        for i, segment in enumerate(text_segments):
            row = table.add_row().cells
            
            p_cues = row[0].paragraphs[0]
            p_cues.add_run(f"[{i+1}]").bold = True
            p_cues.add_run("\nKeywords/Questions:")
            p_cues.add_run(f"\n{segment['translated_text']}").italic = True 
            
            p_notes = row[1].paragraphs[0]
            p_notes.add_run(f"[{i+1}]").bold = True
            p_notes.add_run("\nOriginal Text:")
            p_notes.add_run(f"\n{segment['text']}")
            
            row[1].add_paragraph("\n\nYOUR NOTES HERE: _____________________________________")
            
            try:
                row.height = Inches(3) 
            except Exception:
                pass

        document.add_page_break()
        document.add_heading("Summary Section (for all segments)", level=3)

        summary_table = document.add_table(rows=1, cols=1, style='Table Grid')
        summary_table.autofit = False
        
        # Ширина: 6.5 дюймів
        summary_table.columns[0].width = Inches(6.5) 
        
        cell = summary_table.cell(0, 0)
        cell.text = 'SUMMARY (Висновки та резюме після прослуховування):'
        
        row = summary_table.add_row()
        row.cells[0].text = ''
        try:
            row.height = Inches(1.5)
        except Exception:
            pass
        
        document.add_paragraph("---", style='Normal')


    # --- Функція TTS ---
    def _generate_tts_audio_with_pauses(self, segments, base_name, output_dir, output_format, silence, temp_dir):
        source_lang = self.source_lang.get()
        if source_lang == "auto": source_lang = "en" 
        
        output_format = output_format.lower()
        audio_pauses_tts_filename = os.path.join(output_dir, f"{base_name}_TTS_with_pauses.{output_format}")
        audio_with_pauses_tts = AudioSegment.empty()
        
        self.log_message("   [TTS] Починаю синтез та склеювання...")

        for i, segment in enumerate(segments):
            text = segment['text'].strip()
            if not text: continue
            try:
                tts = gTTS(text=text, lang=source_lang, slow=False)
                mp3_fp = io.BytesIO()
                tts.write_to_fp(mp3_fp)
                mp3_fp.seek(0)
                segment_audio_tts = AudioSegment.from_file(mp3_fp, format="mp3")
                
                if len(segment_audio_tts) > self.FADE_DURATION_MS * 2:
                    segment_audio_tts = segment_audio_tts.fade_in(self.FADE_DURATION_MS).fade_out(self.FADE_DURATION_MS)
                elif len(segment_audio_tts) > 0:
                    segment_audio_tts = segment_audio_tts.fade_in(1).fade_out(1)

                audio_with_pauses_tts += segment_audio_tts
                audio_with_pauses_tts += silence
                
            except Exception as e:
                self.log_message(f"   [TTS] ⚠️ Помилка генерації сегменту {i+1}: {e}. Пропускаю.")
                continue

        if len(audio_with_pauses_tts) > 0:
            # Нормалізація гучності TTS
            change_in_db = -20.0 - audio_with_pauses_tts.dBFS
            audio_with_pauses_tts = audio_with_pauses_tts.apply_gain(change_in_db)
            
            audio_with_pauses_tts.export(audio_pauses_tts_filename, format=output_format)
            self.log_message(f"   [TTS] ✅ Аудіо TTS з паузами збережено: {os.path.basename(audio_pauses_tts_filename)}")
            
            waveform_tts_filename = os.path.join(temp_dir, f"{base_name}_TTS_waveform_with_pauses.png")
            self._plot_waveform(audio_with_pauses_tts, waveform_tts_filename, f"TTS Аудіо з паузами ({len(audio_with_pauses_tts) / 1000:.1f} сек)")
            
            return audio_pauses_tts_filename, waveform_tts_filename
        
        self.log_message("   [TTS] ❌ TTS-версія не була згенерована (ймовірно, помилка gTTS або немає сегментів).")
        return None, None

    def _plot_waveform(self, audio_segment, filename, title):
        samples = np.array(audio_segment.get_array_of_samples())
        if audio_segment.channels == 2:
            samples = samples.reshape((-1, 2))
        
        # Обмеження на кількість точок для швидкого відображення
        max_points = 50000 
        if len(samples) > max_points:
            step = len(samples) // max_points
            samples = samples[::step]
        
        times = np.linspace(0, len(audio_segment) / 1000, len(samples))
        
        plt.figure(figsize=(10, 3))
        if audio_segment.channels == 1:
            plt.plot(times, samples, color='blue')
        else:
            plt.plot(times, samples[:, 0], color='blue', alpha=0.7)
            plt.plot(times, samples[:, 1], color='red', alpha=0.7)

        plt.title(title)
        plt.xlabel("Час (секунди)")
        plt.ylabel("Амплітуда")
        plt.tight_layout()
        plt.savefig(filename)
        plt.close()

    def _process_audio(self):
        audio_filename = self.audio_filepath
        model_name = self.model_name.get()
        pause_seconds = self.pause_duration.get()
        pause_duration_ms = pause_seconds * 1000
        output_dir = self.output_dir.get() 
        source_lang = self.source_lang.get()
        target_lang = self.target_lang.get()
        audio_type = self.audio_type.get()
        generate_tts = self.generate_tts.get()
        segmentation_type = self.segmentation_choice.get()
        note_template = self.note_template.get() 
        output_format = self.output_format.get().lower() 
        
        # --- 1. Підготовка ---
        self.log_message("⏳ Крок 1/5: Завантаження аудіо та моделі Whisper...")
        model_loaded = whisper.load_model(model_name, device=self.device)
        audio_segment = AudioSegment.from_file(audio_filename)
        audio_length_ms = len(audio_segment) 
        frame_rate = audio_segment.frame_rate
        silence = AudioSegment.silent(duration=pause_duration_ms, frame_rate=frame_rate)
        
        temp_dir = tempfile.mkdtemp(prefix=f"{os.path.splitext(os.path.basename(audio_filename))[0]}_whisper_temp_")
        self.log_message(f"✅ Тимчасова папка створена: {temp_dir}")
        base_name = os.path.splitext(os.path.basename(audio_filename))[0]
        
        audio_orig_filename_temp = os.path.join(temp_dir, f"{base_name}_original.mp3")
        audio_pauses_filename_temp = os.path.join(output_dir, f"{base_name}_with_pauses.{output_format}") 
        waveform_orig_filename_temp = os.path.join(temp_dir, f"{base_name}_waveform_original.png")
        waveform_pauses_filename_temp = os.path.join(temp_dir, f"{base_name}_waveform_with_pauses.png")
        zip_filename_final = os.path.join(output_dir, f"{base_name}_material_for_students.zip")
        
        audio_pauses_tts_filename = None
        waveform_tts_filename = None
        audio_segment.export(audio_orig_filename_temp, format="mp3")

        # --- 2/3. Транскрипція та Сегментація ---
        self.log_message("⏳ Крок 2/5: Транскрипція з вирівнюванням за словами...")
        lang_arg = source_lang if source_lang != "auto" else None
        
        result = model_loaded.transcribe(audio_filename, verbose=False, language=lang_arg, word_timestamps=True)
        initial_segments = result.get('segments', [])
        
        if segmentation_type == "Refined (Full Sentences)":
            self.log_message("⏳ Крок 3/5: Фільтрація сегментів за логікою повних речень...")
            semantic_segments = self._refine_segments_by_sentences(initial_segments)
        else:
            self.log_message("⏳ Крок 3/5: Використання базової сегментації Whisper (короткі фрази)...")
            semantic_segments = [
                {'start': s['start'], 'end': s['end'], 'text': s['text'].strip()} 
                for s in initial_segments
            ]
        
        if not semantic_segments: 
            raise ValueError(f"Whisper не зміг розділити аудіо на сегменти. Спробуйте іншу модель або перевірте якість аудіо.")
            
        self.log_message(f"✅ Згенеровано {len(semantic_segments)} сегментів.")

        # --- 4А. Створення аудіо з паузами (ОРИГІНАЛ) та обробка тексту ---
        self.log_message("⏳ Крок 4А/5: Генерація АУДІО З ПАУЗАМИ (ОРИГІНАЛ) та переклад тексту...")
        
        audio_with_pauses = AudioSegment.empty()
        audio_with_pauses = audio_with_pauses.set_frame_rate(frame_rate).set_channels(audio_segment.channels)
        
        doc_info = {
            "Модель Whisper": model_name, 
            "Пауза": f"{pause_seconds} сек.", 
            "Мова тексту": source_lang.upper(), 
            "Мова перекладу (ключ)": target_lang.upper(),
            "Обробка оригіналу": audio_type, 
            "Додаткове TTS-аудіо": 'ТАК' if generate_tts else 'НІ',
            "Логіка сегментації": segmentation_type,
            "Шаблон нотування": note_template
        }
        
        processed_segments = []
        
        for i, segment in enumerate(semantic_segments):
            start_ms = int(segment['start'] * 1000)
            end_ms = int(segment['end'] * 1000)
            text = segment['text'].strip()
            
            try:
                # Використовуємо self.translator (Google Translate)
                translated_text = self.translator.translate(text, dest=target_lang).text
            except Exception as e:
                translated_text = "ПОМИЛКА ПЕРЕКЛАДУ"
                
            segment['translated_text'] = translated_text
            processed_segments.append(segment)
                
            start_buffered = max(0, start_ms - self.AUDIO_START_BUFFER_MS)
            end_buffered = min(audio_length_ms, end_ms + self.AUDIO_END_BUFFER_MS)
            segment_audio = audio_segment[start_buffered:end_buffered]
            
            if audio_type == "TTS (Apply Fade-in/out)":
                if len(segment_audio) > self.FADE_DURATION_MS * 2:
                    segment_audio = segment_audio.fade_in(self.FADE_DURATION_MS).fade_out(self.FADE_DURATION_MS)
                elif len(segment_audio) > 0:
                    segment_audio = segment_audio.fade_in(1).fade_out(1)
            
            if len(segment_audio) < self.MIN_SEGMENT_DURATION_MS:
                continue

            audio_with_pauses += segment_audio
            audio_with_pauses += silence
            
        # Зменшення гучності, щоб уникнути кліппінгу
        change_in_db_orig = -20.0 - audio_with_pauses.dBFS
        audio_with_pauses = audio_with_pauses.apply_gain(change_in_db_orig)

        audio_with_pauses.export(audio_pauses_filename_temp, format=output_format)
        self.log_message(f"✅ Аудіо з паузами (ОРИГІНАЛ) збережено: {os.path.basename(audio_pauses_filename_temp)}")
        
        # --- 4Б. ГЕНЕРАЦІЯ ДОДАТКОВОГО TTS АУДІО ---
        if generate_tts:
             self.log_message("⏳ Крок 4Б/5: Генерація ДОДАТКОВОГО TTS-АУДІО...")
             audio_pauses_tts_filename, waveform_tts_filename = self._generate_tts_audio_with_pauses(
                 processed_segments, base_name, output_dir, output_format, silence, temp_dir
             )
             
        # =================================================================
        # --- 4В. ГЕНЕРАЦІЯ ШАБЛОНІВ НОТУВАННЯ (ОКРЕМІ ФАЙЛИ) ---
        # =================================================================
        self.log_message(f"⏳ Крок 4В/5: Генерація шаблонів нотування '{note_template}'...")

        doc_files_to_zip = []
        
        template_map = {
            "Handout (Текст та Ключ)": (self._create_handout_layout, "_handout", True),
            "Symbolic/Keyword (2 Колонки)": (self._create_symbolic_keyword_layout, "_symbolic_keyword", False),
            "Vertical (3 Колонки / Розан)": (self._create_vertical_layout, "_vertical_rozan", False),
            "Horizontal (Лінійні Нотатки)": (self._create_horizontal_layout, "_horizontal_linear", False),
            "Cornell Notes (Корнельська Система)": (self._create_cornell_layout, "_cornell_notes", False)
        }

        # Helper function to generate a single DOCX
        def generate_doc(layout_func, suffix, segments, target_lang, is_handout):
            doc = Document()
            self._add_document_header(doc, base_name, doc_info)
            if is_handout:
                layout_func(doc, segments) 
            else:
                layout_func(doc, segments, target_lang)
            filename = os.path.join(output_dir, f"{base_name}{suffix}.docx")
            doc.save(filename)
            self.log_message(f"   [DOCX] ✅ Документ Word збережено: {os.path.basename(filename)}")
            return filename
            
        
        if note_template == "Усі 5 Шаблонів (Окремі файли)":
            self.log_message("   [DOCX] Генерую ВСІ 5 шаблонів окремими файлами...")
            
            # List of all templates to generate
            all_templates_config = [
                ("Handout (Текст та Ключ)", "_handout", True),
                ("Symbolic/Keyword (2 Колонки)", "_symbolic_keyword", False),
                ("Vertical (3 Колонки / Розан)", "_vertical_rozan", False),
                ("Horizontal (Лінійні Нотатки)", "_horizontal_linear", False),
                ("Cornell Notes (Корнельська Система)", "_cornell_notes", False)
            ]
            
            for name, suffix, is_handout in all_templates_config:
                func, _, _ = template_map[name]
                doc_filename = generate_doc(func, suffix, processed_segments, target_lang, is_handout)
                doc_files_to_zip.append(doc_filename)
                
        else: # Обробка двох основних шаблонів, вибраних користувачем
            # Знаходимо функцію та суфікс для вибраного шаблону
            template_name_ui = note_template 
            if template_name_ui not in template_map:
                template_name_ui = "Symbolic/Keyword (2 Колонки)" # Fallback

            func, suffix, is_handout = template_map[template_name_ui]
            doc_filename = generate_doc(func, suffix, processed_segments, target_lang, is_handout=is_handout)
            doc_files_to_zip.append(doc_filename)


        # --- 5. Генерація графіків та архівація ---
        self.log_message("⏳ Крок 5/5: Побудова графіків та архівація матеріалів...")
        self._plot_waveform(audio_segment, waveform_orig_filename_temp, f"Оригінальне аудіо ({len(audio_segment) / 1000:.1f} сек)")
        self._plot_waveform(audio_with_pauses, waveform_pauses_filename_temp, f"Аудіо з паузами (Оригінал) ({len(audio_with_pauses) / 1000:.1f} сек)")
        
        
        files_to_zip = [
            audio_orig_filename_temp, 
            audio_pauses_filename_temp, 
            waveform_orig_filename_temp, 
            waveform_pauses_filename_temp
        ]
        
        files_to_zip.extend(doc_files_to_zip) # Додаємо згенеровані DOCX файли
        
        if audio_pauses_tts_filename:
             files_to_zip.append(audio_pauses_tts_filename)
             if waveform_tts_filename:
                 files_to_zip.append(waveform_tts_filename)

        
        self.log_message(f"   [ZIP] Створюю архів {os.path.basename(zip_filename_final)}...")
        
        try:
            with zipfile.ZipFile(zip_filename_final, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for file_path in files_to_zip:
                    # Додаємо файл до ZIP, використовуючи лише його базове ім'я
                    if os.path.exists(file_path):
                        zipf.write(file_path, os.path.basename(file_path))
                    else:
                        self.log_message(f"   [ZIP] ⚠️ Файл не знайдено для архівації: {os.path.basename(file_path)}")
            self.log_message(f"   [ZIP] ✅ Усі матеріали заархівовано до: {zip_filename_final}")
        except Exception as e:
            self.log_message(f"   [ZIP] ❌ Помилка при створенні архіву: {e}")
            
        return temp_dir

    def _cleanup(self, temp_dir):
        if os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
                self.log_message(f"✅ Тимчасова папка видалена.")
            except Exception as e:
                self.log_message(f"❌ Помилка видалення тимчасової папки: {e}")


if __name__ == "__main__":
    app = ConsecutiveTrainerApp()
    app.mainloop()