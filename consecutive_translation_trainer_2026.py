# -*- coding: utf-8 -*-
"""
Consecutive Translation Trainer (Complete Local Version)
Повна виправлена та оптимізована версія з динамічним пошуком шляхів,
багатопотоковим розпізнаванням Whisper, перекладом та експортом в DOCX.
"""

import sys
import os
import time
import shutil
import threading
import pathlib
import tempfile
import re
import io
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext

# Налаштування кешу Whisper у безпечній папці користувача
WHISPER_CACHE_ROOT = os.path.join(os.path.expanduser("~"), "Whisper_cache")
os.makedirs(WHISPER_CACHE_ROOT, exist_ok=True)
os.environ["XDG_CACHE_HOME"] = WHISPER_CACHE_ROOT

# Додавання можливих системних шляхів до site-packages якщо вони відсутні в sys.path
possible_site_packages = [
    r"C:\Users\Tata\AppData\Local\Programs\Python\Python311\Lib\site-packages",
    os.path.join(os.path.expanduser("~"), "AppData", "Local", "Programs", "Python", "Python311", "Lib", "site-packages")
]
for sp in possible_site_packages:
    if os.path.exists(sp) and sp not in sys.path:
        sys.path.append(sp)

def find_ffmpeg():
    if shutil.which("ffmpeg"):
        return True
    
    common_paths = [
        r"C:\ffmpeg\bin",
        r"C:\Program Files\ffmpeg\bin",
        r"C:\Users\Tata\ffmpeg-2025-09-15-git-16b8a7805b-full_build\bin",
        os.path.join(os.path.expanduser("~"), "ffmpeg", "bin")
    ]
    for path in common_paths:
        if os.path.exists(os.path.join(path, "ffmpeg.exe")):
            os.environ["PATH"] += os.pathsep + path
            return True
    return False

if find_ffmpeg():
    print("✅ [ENV] FFMPEG успішно знайдено та підключено.")
else:
    print("⚠️ [ENV] Попередження: FFMPEG не знайдено автоматично. Переконайтеся, що він доданий до PATH.")

try:
    import torch
    import numpy as np
    import whisper
    from pydub import AudioSegment
    from docx import Document
    from docx.shared import Pt, Inches
    from deep_translator import GoogleTranslator
    from gtts import gTTS
    import customtkinter as ctk
except ImportError as e:
    print(f"❌ КРИТИЧНА ПОМИЛКА імпорту: {e}")
    print(">>> Будь ласка, встановіть необхідні бібліотеки командою:")
    print("pip install torch numpy openai-whisper pydub python-docx deep-translator gtts customtkinter")
    sys.exit(1)

DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
print(f"🚀 Використовуваний пристрій для обробки: {DEVICE}")

class TranslationTrainerApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.title("Consecutive Translation Trainer (Pro)")
        self.geometry("1050x750")
        
        # Налаштування теми CustomTkinter
        ctk.set_appearance_mode("System")
        ctk.set_default_color_theme("blue")

        self.audio_file_path = None
        self.whisper_model = None

        self.create_widgets()

    def create_widgets(self):
        self.grid_rowconfigure(1, weight=1)
        self.grid_columnconfigure(0, weight=1)

        # --- Верхня панель керування ---
        self.top_frame = ctk.CTkFrame(self, height=70)
        self.top_frame.grid(row=0, column=0, sticky="ew", padx=10, pady=10)

        self.btn_load = ctk.CTkButton(self.top_frame, text="📁 Завантажити аудіо", command=self.load_audio_file, width=150)
        self.btn_load.pack(side=tk.LEFT, padx=10, pady=15)

        self.model_label = ctk.CTkLabel(self.top_frame, text="Whisper модель:")
        self.model_label.pack(side=tk.LEFT, padx=(10, 0), pady=15)

        self.model_var = ctk.StringVar(value="base")
        self.model_menu = ctk.CTkOptionMenu(self.top_frame, values=["tiny", "base", "small", "medium"], variable=self.model_var, width=100)
        self.model_menu.pack(side=tk.LEFT, padx=5, pady=15)

        self.btn_transcribe = ctk.CTkButton(self.top_frame, text="🎙️ Розпізнати (Whisper)", command=self.start_transcription_thread, fg_color="#2b8a3e", hover_color="#237032", width=170)
        self.btn_transcribe.pack(side=tk.LEFT, padx=10, pady=15)

        self.lbl_status = ctk.CTkLabel(self.top_frame, text="Статус: Очікування файлу...", text_color="orange")
        self.lbl_status.pack(side=tk.LEFT, padx=15, pady=15)

        # --- Центральна робоча зона (Текстові поля) ---
        self.center_frame = ctk.CTkFrame(self)
        self.center_frame.grid(row=1, column=0, sticky="nsew", padx=10, pady=5)
        self.center_frame.grid_rowconfigure(1, weight=1)
        self.center_frame.grid_columnconfigure((0, 1), weight=1)

        # Колонка оригінального тексту
        self.lbl_orig = ctk.CTkLabel(self.center_frame, text="Оригінальний текст (Розпізнаний):", font=("Arial", 13, "bold"))
        self.lbl_orig.grid(row=0, column=0, sticky="w", padx=5, pady=5)

        self.txt_orig = scrolledtext.ScrolledText(self.center_frame, wrap=tk.WORD, width=45, height=22, font=("Arial", 11))
        self.txt_orig.grid(row=1, column=0, sticky="nsew", padx=5, pady=5)

        # Колонка перекладу
        self.lbl_trans = ctk.CTkLabel(self.center_frame, text="Переклад (Цільова мова):", font=("Arial", 13, "bold"))
        self.lbl_trans.grid(row=0, column=1, sticky="w", padx=5, pady=5)

        self.txt_trans = scrolledtext.ScrolledText(self.center_frame, wrap=tk.WORD, width=45, height=22, font=("Arial", 11))
        self.txt_trans.grid(row=1, column=1, sticky="nsew", padx=5, pady=5)

        # --- Нижня панель дій ---
        self.bottom_frame = ctk.CTkFrame(self, height=60)
        self.bottom_frame.grid(row=2, column=0, sticky="ew", padx=10, pady=10)

        self.lang_var = ctk.StringVar(value="uk")
        self.lang_menu = ctk.CTkOptionMenu(self.bottom_frame, values=["uk", "en", "pl", "de", "fr", "es"], variable=self.lang_var, width=80)
        self.lang_menu.pack(side=tk.LEFT, padx=10, pady=10)

        self.btn_translate = ctk.CTkButton(self.bottom_frame, text="🌐 Перекласти текст", command=self.translate_text, width=150)
        self.btn_translate.pack(side=tk.LEFT, padx=5, pady=10)

        self.btn_export = ctk.CTkButton(self.bottom_frame, text="📄 Експорт в Word (.docx)", command=self.export_to_docx, fg_color="#1864ab", hover_color="#104e8b", width=180)
        self.btn_export.pack(side=tk.LEFT, padx=15, pady=10)

    def load_audio_file(self):
        file_path = filedialog.askopenfilename(
            title="Виберіть аудіофайл для тренування",
            filetypes=[("Audio files", "*.mp3 *.wav *.m4a *.flac *.ogg"), ("All files", "*.*")]
        )
        if file_path:
            self.audio_file_path = file_path
            filename = os.path.basename(file_path)
            self.lbl_status.configure(text=f"Завантажено: {filename}", text_color="green")
            self.txt_orig.delete("1.0", tk.END)
            self.txt_orig.insert(tk.END, f"Файл готовий до обробки:\n{file_path}\n\nНатисніть кнопку 'Розпізнати (Whisper)' для початку.")

    def start_transcription_thread(self):
        if not self.audio_file_path or not os.path.exists(self.audio_file_path):
            messagebox.showwarning("Попередження", "Будь ласка, спочатку завантажте дійсний аудіофайл!")
            return
        
        self.btn_transcribe.configure(state="disabled")
        self.btn_load.configure(state="disabled")
        self.lbl_status.configure(text="Статус: Завантаження моделі та розпізнавання...", text_color="blue")
        
        thread = threading.Thread(target=self.run_whisper_transcription)
        thread.daemon = True
        thread.start()

    def run_whisper_transcription(self):
        try:
            model_name = self.model_var.get()
            print(f"📦 Завантаження моделі Whisper ({model_name})... на пристрої {DEVICE}")
            model = whisper.load_model(model_name, device=DEVICE)
            
            print(f"🎙️ Початок розпізнавання файлу: {self.audio_file_path}")
            result = model.transcribe(self.audio_file_path)
            transcript_text = result.get("text", "").strip()

            # Оновлення UI у головному потоці
            self.after(0, lambda: self.finish_transcription(transcript_text))
        except Exception as e:
            error_msg = str(e)
            print(f"❌ Помилка розпізнавання: {error_msg}")
            self.after(0, lambda: self.show_error(error_msg))

    def finish_transcription(self, text):
        self.txt_orig.delete("1.0", tk.END)
        self.txt_orig.insert(tk.END, text)
        self.lbl_status.configure(text="Статус: Розпізнавання успішно завершено!", text_color="green")
        self.btn_transcribe.configure(state="normal")
        self.btn_load.configure(state="normal")

    def show_error(self, error_msg):
        messagebox.showerror("Помилка Whisper", f"Сталася помилка під час розпізнавання:\n{error_msg}")
        self.lbl_status.configure(text="Статус: Помилка обробки", text_color="red")
        self.btn_transcribe.configure(state="normal")
        self.btn_load.configure(state="normal")

    def translate_text(self):
        text_to_translate = self.txt_orig.get("1.0", tk.END).strip()
        if not text_to_translate or "Файл готовий до обробки" in text_to_translate:
            messagebox.showwarning("Попередження", "Немає розпізнаного тексту для перекладу!")
            return
        
        target_lang = self.lang_var.get()
        try:
            self.lbl_status.configure(text="Статус: Виконується переклад...", text_color="blue")
            
            # Переклад шматками, якщо текст дуже довгий (deep-translator має ліміти за символами)
            chunks = [text_to_translate[i:i+4000] for i in range(0, len(text_to_translate), 4000)]
            translated_full = ""
            
            for chunk in chunks:
                translated_chunk = GoogleTranslator(source='auto', target=target_lang).translate(chunk)
                translated_full += translated_chunk + "\n"

            self.txt_trans.delete("1.0", tk.END)
            self.txt_trans.insert(tk.END, translated_full.strip())
            self.lbl_status.configure(text="Статус: Переклад успішно завершено!", text_color="green")
        except Exception as e:
            messagebox.showerror("Помилка перекладу", f"Не вдалося виконати переклад: {e}")
            self.lbl_status.configure(text="Статус: Помилка перекладу", text_color="red")

    def export_to_docx(self):
        try:
            orig_text = self.txt_orig.get("1.0", tk.END).strip()
            trans_text = self.txt_trans.get("1.0", tk.END).strip()
            
            if not orig_text:
                messagebox.showwarning("Попередження", "Немає даних для експорту!")
                return

            doc = Document()
            doc.add_heading('Consecutive Translation Training Report', 0)
            
            doc.add_heading('Original Transcript / Text:', level=1)
            doc.add_paragraph(orig_text)
            
            doc.add_heading('Translation:', level=1)
            doc.add_paragraph(trans_text if trans_text else "[Переклад відсутній]")
            
            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Document", "*.docx")],
                title="Зберегти звіт у Word"
            )
            if save_path:
                doc.save(save_path)
                messagebox.showinfo("Успіх", f"Документ успішно збережено:\n{save_path}")
                self.lbl_status.configure(text="Статус: Експорт у Word завершено", text_color="green")
        except Exception as e:
            messagebox.showerror("Помилка експорту", f"Не вдалося зберегти файл: {e}")

if __name__ == "__main__":
    app = TranslationTrainerApp()
    app.mainloop()